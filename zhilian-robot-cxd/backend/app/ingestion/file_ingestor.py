"""
文件接入器 - 支持PDF、Excel、CSV、Word等格式
"""
import io
import os
import logging
from typing import List, Dict, Any, Optional, BinaryIO
from datetime import datetime

from .base_ingestor import BaseIngestor, RawDataRecord

logger = logging.getLogger(__name__)


class FileIngestor(BaseIngestor):
    """
    通用文件接入器
    
    根据文件类型自动选择合适的解析方法
    """
    
    # 支持的文件类型映射
    SUPPORTED_TYPES = {
        '.pdf': 'application/pdf',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.html': 'text/html',
    }
    
    def __init__(self, source_name: str = "file_upload"):
        super().__init__(source_type="file", source_name=source_name)
    
    def ingest(self, source: Any, **kwargs) -> List[RawDataRecord]:
        """
        接入文件数据
        
        Args:
            source: 可以是文件路径(str)或文件对象(BinaryIO)
            **kwargs:
                - filename: 原始文件名 (当source是BinaryIO时需要)
                - metadata: 额外的元数据
                
        Returns:
            原始数据记录列表
        """
        records = []
        
        try:
            # 处理不同类型的输入
            if isinstance(source, str):
                # 文件路径
                if os.path.isfile(source):
                    records.append(self._ingest_file_path(source, **kwargs))
                elif os.path.isdir(source):
                    # 目录: 递归处理所有文件
                    for root, _, files in os.walk(source):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                record = self._ingest_file_path(file_path, **kwargs)
                                records.append(record)
                            except Exception as e:
                                self.logger.warning(f"处理文件失败 {file_path}: {e}")
                else:
                    raise ValueError(f"文件或目录不存在: {source}")
            elif hasattr(source, 'read'):
                # 文件对象
                filename = kwargs.pop('filename', 'unknown')
                records.append(self._ingest_file_object(source, filename, **kwargs))
            elif isinstance(source, bytes):
                # 字节数据
                filename = kwargs.pop('filename', 'unknown')
                records.append(self._ingest_bytes(source, filename, **kwargs))
            else:
                raise ValueError(f"不支持的数据源类型: {type(source)}")
                
        except Exception as e:
            self.logger.error(f"文件接入失败: {e}")
            raise
        
        return records
    
    def _ingest_file_path(self, file_path: str, **kwargs) -> RawDataRecord:
        """从文件路径接入"""
        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()
        content_type = self.SUPPORTED_TYPES.get(ext, 'application/octet-stream')
        
        with open(file_path, 'rb') as f:
            raw_content = f.read()
        
        # 获取文件元数据
        file_stat = os.stat(file_path)
        metadata = kwargs.get('metadata', {})
        metadata.update({
            'filename': filename,
            'file_path': file_path,
            'file_size': file_stat.st_size,
            'file_extension': ext,
            'modified_time': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
        })
        
        return RawDataRecord(
            source_type="file",
            source_name=self.source_name,
            source_url=f"file://{file_path}",
            raw_content=raw_content,
            content_type=content_type,
            metadata=metadata
        )
    
    def _ingest_file_object(self, file_obj: BinaryIO, filename: str, **kwargs) -> RawDataRecord:
        """从文件对象接入"""
        raw_content = file_obj.read()
        ext = os.path.splitext(filename)[1].lower()
        content_type = self.SUPPORTED_TYPES.get(ext, 'application/octet-stream')
        
        metadata = kwargs.get('metadata', {})
        metadata.update({
            'filename': filename,
            'file_extension': ext,
            'file_size': len(raw_content)
        })
        
        return RawDataRecord(
            source_type="file",
            source_name=self.source_name,
            raw_content=raw_content,
            content_type=content_type,
            metadata=metadata
        )
    
    def _ingest_bytes(self, data: bytes, filename: str, **kwargs) -> RawDataRecord:
        """从字节数据接入"""
        ext = os.path.splitext(filename)[1].lower()
        content_type = self.SUPPORTED_TYPES.get(ext, 'application/octet-stream')
        
        metadata = kwargs.get('metadata', {})
        metadata.update({
            'filename': filename,
            'file_extension': ext,
            'file_size': len(data)
        })
        
        return RawDataRecord(
            source_type="file",
            source_name=self.source_name,
            raw_content=data,
            content_type=content_type,
            metadata=metadata
        )
    
    def validate(self, record: RawDataRecord) -> bool:
        """验证文件记录"""
        # 检查内容不为空
        if not record.raw_content:
            self.logger.warning("文件内容为空")
            return False
        
        # 检查文件大小 (限制100MB)
        max_size = 100 * 1024 * 1024
        if len(record.raw_content) > max_size:
            self.logger.warning(f"文件过大: {len(record.raw_content)} bytes (最大 {max_size})")
            return False
        
        return True
    
    def extract_text(self, record: RawDataRecord) -> str:
        """
        从文件中提取文本
        
        根据content_type选择合适的提取方法
        """
        content_type = record.content_type
        
        try:
            if content_type == 'application/pdf':
                return self._extract_pdf_text(record.raw_content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                  'application/vnd.ms-excel']:
                return self._extract_excel_text(record.raw_content)
            elif content_type == 'text/csv':
                return self._extract_csv_text(record.raw_content, record.content_encoding)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                  'application/msword']:
                return self._extract_docx_text(record.raw_content)
            elif content_type.startswith('text/'):
                return record.raw_content.decode(record.content_encoding)
            else:
                self.logger.warning(f"不支持提取文本的类型: {content_type}")
                return ""
        except Exception as e:
            self.logger.error(f"文本提取失败: {e}")
            return ""
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """提取PDF文本"""
        try:
            import PyPDF2
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            self.logger.error("PyPDF2未安装，无法提取PDF文本")
            return ""
    
    def _extract_excel_text(self, content: bytes) -> str:
        """提取Excel文本"""
        try:
            import pandas as pd
            excel_file = io.BytesIO(content)
            
            # 读取所有sheet
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_parts = []
            for sheet_name, df in excel_data.items():
                text_parts.append(f"=== {sheet_name} ===")
                # 转换为字符串表格
                text_parts.append(df.to_string())
            
            return "\n\n".join(text_parts)
        except ImportError:
            self.logger.error("pandas/openpyxl未安装，无法提取Excel文本")
            return ""
    
    def _extract_csv_text(self, content: bytes, encoding: str = 'utf-8') -> str:
        """提取CSV文本"""
        try:
            import pandas as pd
            
            # 尝试检测编码
            try:
                import chardet
                detected = chardet.detect(content)
                encoding = detected.get('encoding', encoding) or encoding
            except ImportError:
                pass
            
            csv_file = io.BytesIO(content)
            df = pd.read_csv(csv_file, encoding=encoding)
            return df.to_string()
        except ImportError:
            self.logger.error("pandas未安装，无法提取CSV文本")
            return ""
        except Exception as e:
            self.logger.error(f"CSV解析失败: {e}")
            return content.decode(encoding, errors='ignore')
    
    def _extract_docx_text(self, content: bytes) -> str:
        """提取Word文档文本"""
        try:
            from docx import Document
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except ImportError:
            self.logger.error("python-docx未安装，无法提取Word文本")
            return ""


class PDFIngestor(FileIngestor):
    """PDF文件专用接入器"""
    
    def __init__(self, source_name: str = "pdf_upload"):
        super().__init__(source_name=source_name)
    
    def validate(self, record: RawDataRecord) -> bool:
        """验证PDF文件"""
        if not super().validate(record):
            return False
        
        # 检查PDF魔数
        if not record.raw_content.startswith(b'%PDF'):
            self.logger.warning("无效的PDF文件格式")
            return False
        
        return True
    
    def extract_tables(self, record: RawDataRecord) -> List[Dict]:
        """
        提取PDF中的表格数据
        
        Returns:
            表格列表，每个表格是一个字典列表
        """
        try:
            import pdfplumber
            pdf_file = io.BytesIO(record.raw_content)
            
            tables = []
            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    for table_num, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                'page': page_num + 1,
                                'table_index': table_num,
                                'data': table
                            })
            
            return tables
        except ImportError:
            self.logger.error("pdfplumber未安装，无法提取PDF表格")
            return []


class ExcelIngestor(FileIngestor):
    """Excel文件专用接入器"""
    
    def __init__(self, source_name: str = "excel_upload"):
        super().__init__(source_name=source_name)
    
    def validate(self, record: RawDataRecord) -> bool:
        """验证Excel文件"""
        if not super().validate(record):
            return False
        
        # 检查Excel魔数 (xlsx是zip格式)
        if record.raw_content.startswith(b'PK'):
            return True
        # 老版本xls格式
        if record.raw_content.startswith(b'\xd0\xcf\x11\xe0'):
            return True
        
        self.logger.warning("无效的Excel文件格式")
        return False
    
    def to_dataframe(self, record: RawDataRecord, sheet_name: str = None):
        """
        将Excel转换为DataFrame
        
        Args:
            record: 原始数据记录
            sheet_name: 指定sheet名称，None表示所有sheet
            
        Returns:
            DataFrame或DataFrame字典
        """
        try:
            import pandas as pd
            excel_file = io.BytesIO(record.raw_content)
            return pd.read_excel(excel_file, sheet_name=sheet_name)
        except ImportError:
            self.logger.error("pandas未安装")
            return None


class CSVIngestor(FileIngestor):
    """CSV文件专用接入器"""
    
    def __init__(self, source_name: str = "csv_upload"):
        super().__init__(source_name=source_name)
    
    def to_dataframe(self, record: RawDataRecord, **pandas_kwargs):
        """
        将CSV转换为DataFrame
        
        Args:
            record: 原始数据记录
            **pandas_kwargs: 传递给pd.read_csv的参数
            
        Returns:
            DataFrame
        """
        try:
            import pandas as pd
            
            # 检测编码
            encoding = record.content_encoding
            try:
                import chardet
                detected = chardet.detect(record.raw_content)
                encoding = detected.get('encoding', encoding) or encoding
            except ImportError:
                pass
            
            csv_file = io.BytesIO(record.raw_content)
            return pd.read_csv(csv_file, encoding=encoding, **pandas_kwargs)
        except ImportError:
            self.logger.error("pandas未安装")
            return None


# 便捷函数
def ingest_file(file_path: str, source_name: str = "file_upload") -> List[RawDataRecord]:
    """
    便捷函数: 接入单个文件
    
    Args:
        file_path: 文件路径
        source_name: 来源名称
        
    Returns:
        原始数据记录列表
    """
    ingestor = FileIngestor(source_name=source_name)
    return ingestor.ingest(file_path)


def ingest_and_store_file(file_path: str, source_name: str = "file_upload") -> List[RawDataRecord]:
    """
    便捷函数: 接入文件并存储到MinIO和MongoDB
    
    Args:
        file_path: 文件路径
        source_name: 来源名称
        
    Returns:
        处理后的记录列表
    """
    ingestor = FileIngestor(source_name=source_name)
    return ingestor.ingest_and_store(file_path)
